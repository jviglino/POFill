import docx
from tkinter import *
from datetime import *

now = datetime.now()

# Main Window:
window = Tk()
window.title('LinguaPros PO Creation')
window.geometry('300x400')

VendorName = Label(text="Vendor name:")
VendorName.grid(row=1, column=0, sticky=W)
PM_label = Label(text="PM name:")
PM_label.grid(row=3, column=0, sticky=W)
LP_label = Label(text="Language Pair:")
LP_label.grid(row=5, column=0, sticky=W)
Ref_label = Label(text="Job Number:")
Ref_label.grid(row=1, column=4, sticky=W)
Duedate_label = Label(text="Due Date:")
Duedate_label.grid(row=3, column=4, sticky=W)
Jobdescription_label = Label(text="Job Type:")
Jobdescription_label.grid(row=5, column=4, sticky=W)
Notes_label = Label(text="Description:")
Notes_label.grid(row=7, column=0, sticky=W)

#Entries
name_entry = Entry(window, width=20)
name_entry.grid(row=2, column=0, sticky=W)
PM_entry = Entry(window, width=20)
PM_entry.grid(row=4, column=0, sticky=W)
LP_entry = Entry(window, width=20)
LP_entry.grid(row=6, column=0, sticky=W)
Ref_entry = Entry(window, width=20)
Ref_entry.grid(row=2, column=4, sticky=W)
DDvar = StringVar(window, value='mm/dd/yyyy')
Duedate_entry = Entry(window, textvariable=DDvar,  width=20)
Duedate_entry.grid(row=4, column=4, sticky=W)
jobdvar = StringVar(window)
jobdvar.set("Translation")
Jobdescription = OptionMenu(window, jobdvar, "Translation", "Editing", "DTP", "Transcription")
Jobdescription.grid(row=6, column=4, sticky=W)
Notes_entry = Entry(window, width=20)
Notes_entry.grid(row=8, column=0, sticky=W)


# Fill Purchase Order
def pofill():
    global name
    global PM
    global LP
    global Order
    global Ref
    global Due
    global JobD
    global Notes
    name = name_entry.get()
    PM = PM_entry.get()
    LP = LP_entry.get()
    Order = now.strftime("%m/%d/%Y")
    Ref = Ref_entry.get()
    Due = Duedate_entry.get()
    doc = docx.Document('PO Template.docx')
    JobD = jobdvar.get()
    Notes = Notes_entry.get()
    table = doc.tables[0]
    row = table.rows[0]
    row.cells[0].text = 'Resource Name: ' + name
    row.cells[1].text = 'Project Manager: ' + PM
    row = table.rows[1]
    row.cells[0].text = 'Language Pair: ' + LP
    row.cells[1].text = 'Order Date: ' + Order
    row = table.rows[2]
    row.cells[0].text = 'PO Reference: ' + Ref
    row.cells[1].text = 'Due Date: ' + Due
    table = doc.tables[1]
    row = table.rows[1]
    row.cells[0].text = JobD
    row = table.rows[1]
    row.cells[1].text = Notes
    doc.save(Ref + '.docx')

Button(window, text="Create", width=6, command=pofill).grid(row=13, column=2, sticky=W)

window.mainloop()